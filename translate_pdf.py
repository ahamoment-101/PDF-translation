import fitz
from transformers import MarianMTModel, MarianTokenizer
import torch
import time
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import hashlib
import re

def get_image_key(img_bytes):
    """计算图片数据的 MD5 哈希作为唯一标识"""
    m = hashlib.md5()
    m.update(img_bytes)
    return m.hexdigest()

def translate_text(text, tokenizer, model, device, page_num):
    """翻译文本，最多尝试 3 次，返回翻译后的文本（如果中文比例满足要求）"""
    max_attempts = 3
    translated_text = None
    for attempt in range(max_attempts):
        try:
            processed_text = text.replace('\n', ' ').strip()
            print(f"页{page_num+1} 正在翻译文本（前100字符）：{processed_text[:100]}...")
            inputs = tokenizer(
                processed_text,
                return_tensors="pt",
                padding=True,
                truncation=True,
                max_length=256
            ).to(device)
            translated = model.generate(
                **inputs,
                max_length=256,  # 控制输出长度，数值越大，输出越长
                num_beams=25,    # 控制搜索宽度，数值越大，搜索宽度越大
                length_penalty=2.0,  # 控制长度惩罚，数值越大，惩罚越大
                temperature=0.4,     # 控制输出多样性，数值越大，多样性越高
                do_sample=False,     # 是否进行采样
                top_k=50,           # 对输出结果进行截断，数值越大，截断越严格
                top_p=0.6,          # 对输出结果的多样性进行惩罚，数值越大，惩罚越大
                repetition_penalty=2.5  # 对内容重复的惩罚，数值越大，惩罚越大
            )
            candidate_text = tokenizer.batch_decode(translated, skip_special_tokens=True)[0].strip()
            # 计算翻译结果中中文字符比例
            chinese_char_count = sum(1 for char in candidate_text if '\u4e00' <= char <= '\u9fff')
            total_char_count = len(candidate_text)
            chinese_ratio = chinese_char_count / total_char_count if total_char_count > 0 else 0
            print(f"页{page_num+1} 翻译结果（前100字符）：{candidate_text[:100]}...，中文比例：{chinese_ratio:.2%}")
            if chinese_ratio >= 0.15:
                translated_text = candidate_text
                break
            else:
                print(f"页{page_num+1} 第{attempt+1}次翻译中文比例不足，重试...")
                time.sleep(0.5)
        except Exception as e:
            print(f"页{page_num+1} 第{attempt+1}次翻译出错：{e}")
            if attempt == max_attempts - 1:
                translated_text = None
    return translated_text

def join_spans(spans, gap_threshold=3):
    """
    根据 span 的位置还原文本内容：
    如果相邻两个 span 的起始位置间隔大于 gap_threshold，则插入空格
    """
    if not spans:
        return ""
    # 按横坐标排序
    sorted_spans = sorted(spans, key=lambda s: s["bbox"][0])
    text = sorted_spans[0]["text"]
    last_right = sorted_spans[0]["bbox"][2]
    for span in sorted_spans[1:]:
        gap = span["bbox"][0] - last_right
        if gap > gap_threshold:
            text += " " + span["text"]
        else:
            text += span["text"]
        last_right = span["bbox"][2]
    return text

def merge_spans_for_table(spans, gap_threshold=10):
    """
    对于表格识别时，合并同一行中横坐标间隙小于 gap_threshold 的相邻 span，返回一个单元格文本列表
    """
    sorted_spans = sorted(spans, key=lambda s: s["bbox"][0])
    if not sorted_spans:
        return []
    cells = []
    current_cell_text = sorted_spans[0]["text"]
    current_cell_end = sorted_spans[0]["bbox"][2]
    for span in sorted_spans[1:]:
        gap = span["bbox"][0] - current_cell_end
        if gap < gap_threshold:
            current_cell_text += " " + span["text"]
            current_cell_end = span["bbox"][2]
        else:
            cells.append(current_cell_text.strip())
            current_cell_text = span["text"]
            current_cell_end = span["bbox"][2]
    cells.append(current_cell_text.strip())
    return cells

def is_table_block(block, gap_threshold=10):
    """
    判断一个文本块是否为表格候选块：
    - 块内至少有两行
    - 每一行经过 merge_spans_for_table 合并后均至少有两个单元格
    - 且各行的单元格数相同
    """
    lines = block.get("lines", [])
    if len(lines) < 2:
        return False
    cell_counts = []
    for line in lines:
        spans = line.get("spans", [])
        cells = merge_spans_for_table(spans, gap_threshold)
        if len(cells) < 2:
            return False
        cell_counts.append(len(cells))
    return len(set(cell_counts)) == 1  # 每行单元格数一致

def translate_pdf_to_word_with_styles_improved(input_file, output_file):
    try:
        pdf_doc = fitz.open(input_file)
    except Exception as e:
        print(f"打开 PDF 文件失败：{e}")
        return

    # 加载翻译模型和分词器
    model_name = 'Helsinki-NLP/opus-mt-en-zh'
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    model = model.to(device)
    model.eval()

    word_doc = Document()
    # 设置全局文档样式（如果需要，可以进一步调整）
    normal_style = word_doc.styles['Normal']
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.line_spacing = 1.0

    max_width = int(word_doc.sections[0].page_width * 0.7)
    processed_keys = set()

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        print(f"正在处理第 {page_num+1} 页...")
        word_doc.add_heading(f"Page {page_num+1}", level=1)

        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])
        blocks.sort(key=lambda b: b["bbox"][1])

        # 提取本页文本块最左边界，用于缩进参考
        all_x = []
        for block in blocks:
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if spans:
                        all_x.append(spans[0]["bbox"][0])
        page_left_margin = min(all_x) if all_x else 0

        for block in blocks:
            block_type = block.get("type", None)
            if block_type == 0:
                # 表格处理优先
                if is_table_block(block):
                    print(f"页{page_num+1} 检测到表格块，进行表格处理")
                    lines = block.get("lines", [])
                    table_data = []
                    for line in lines:
                        spans = line.get("spans", [])
                        cells = merge_spans_for_table(spans)
                        table_data.append(cells)
                    num_rows = len(table_data)
                    num_cols = len(table_data[0]) if table_data else 0
                    if num_cols == 0:
                        continue
                    table = word_doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = "Table Grid"
                    for i, row in enumerate(table_data):
                        for j in range(num_cols):
                            cell_text = row[j] if j < len(row) else ""
                            # 翻译单元格内容
                            translated_cell = translate_text(cell_text, tokenizer, model, device, page_num)
                            final_text = translated_cell if (translated_cell and translated_cell != cell_text) else cell_text
                            table.cell(i, j).text = final_text
                else:
                    # 对于普通文本块，先将每行的 span 合并，还原原文内容
                    lines_info = []  # 保存 (文本, 首个 span 的起始横坐标)
                    for line in block.get("lines", []):
                        spans = line.get("spans", [])
                        line_text = join_spans(spans)
                        if line_text.strip():
                            first_x = spans[0]["bbox"][0] if spans else page_left_margin
                            lines_info.append((line_text, first_x))
                    if not lines_info:
                        continue

                    # 判断是否为列表：若每行都以常见列表标记开头，则视为列表
                    list_pattern_number = r"^\s*\d+[\.\)]\s+"
                    list_pattern_bullet = r"^\s*[•\-\*\u2022]\s+"
                    is_list = all(re.match(list_pattern_number, text) or re.match(list_pattern_bullet, text)
                                  for text, _ in lines_info)

                    if is_list:
                        for text, first_x in lines_info:
                            style = None
                            if re.match(list_pattern_number, text):
                                style = "List Number"
                                text = re.sub(list_pattern_number, "", text, count=1).strip()
                            elif re.match(list_pattern_bullet, text):
                                style = "List Bullet"
                                text = re.sub(list_pattern_bullet, "", text, count=1).strip()
                            # 翻译处理
                            translated = translate_text(text, tokenizer, model, device, page_num)
                            final_text = translated if (translated and translated != text) else text
                            p = word_doc.add_paragraph(final_text, style=style)
                            # 设置段落格式：缩进与紧凑排版
                            p.paragraph_format.left_indent = Pt(first_x - page_left_margin)
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing = 1.0
                    else:
                        # 非列表，则将所有行合并为一个段落，保留换行
                        combined_text = "\n".join(text for text, _ in lines_info)
                        # 翻译整个块，若翻译后与原文有较大不同，则用翻译结果，否则保留原文
                        translated_block = translate_text(combined_text, tokenizer, model, device, page_num)
                        final_text = translated_block if (translated_block and translated_block != combined_text) else combined_text
                        p = word_doc.add_paragraph(final_text)
                        # 对整个段落设置紧凑格式
                        p.paragraph_format.left_indent = Pt(lines_info[0][1] - page_left_margin)
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.0
            elif block_type == 1:
                # 图片处理逻辑保持不变
                xref = block.get("xref")
                if xref is None:
                    xref = block.get("image")
                if xref is not None:
                    try:
                        if isinstance(xref, bytes):
                            img_bytes = xref
                            img_ext = "png"
                        else:
                            xref = int(xref)
                            img_dict = pdf_doc.extract_image(xref)
                            img_bytes = img_dict.get("image", None)
                            img_ext = img_dict.get("ext", "png")
                        if not img_bytes:
                            print(f"页{page_num+1} 提取图片失败，无法获取图片数据 (xref: {xref})")
                            continue
                        key = get_image_key(img_bytes)
                        if key in processed_keys:
                            print(f"页{page_num+1} 图片重复，跳过")
                            continue
                        processed_keys.add(key)
                        image_stream = BytesIO(img_bytes)
                        image_stream.seek(0)
                        print(f"页{page_num+1} 添加图片（key: {key}, ext: {img_ext}），宽度设置为 {max_width} twips")
                        word_doc.add_picture(image_stream, width=max_width)
                    except Exception as e:
                        print(f"页{page_num+1} 提取图片失败（xref: {xref}）：{e}")
                else:
                    print(f"页{page_num+1} 未在图片块中找到 xref 信息。")
            else:
                print(f"页{page_num+1} 遇到未知块类型：{block_type}")

        # 补充：使用 fallback 方法提取本页中未在 blocks 中识别的图片
        images = page.get_images(full=True)
        for img in images:
            xref = img[0]
            try:
                img_dict = pdf_doc.extract_image(xref)
                img_bytes = img_dict.get("image", None)
                if not img_bytes:
                    print(f"页{page_num+1} fallback 提取图片失败，无法获取图片数据 (xref: {xref})")
                    continue
                key = get_image_key(img_bytes)
                if key in processed_keys:
                    print(f"页{page_num+1} fallback 图片重复，跳过")
                    continue
                processed_keys.add(key)
                img_ext = img_dict.get("ext", "png")
                image_stream = BytesIO(img_bytes)
                image_stream.seek(0)
                print(f"页{page_num+1} fallback 添加图片（key: {key}, ext: {img_ext}），宽度设置为 {max_width} twips")
                word_doc.add_picture(image_stream, width=max_width)
            except Exception as e:
                print(f"页{page_num+1} fallback 提取图片失败（xref: {xref}）：{e}")

        word_doc.add_page_break()

    try:
        word_doc.save(output_file)
        print(f"翻译完成！已保存到 {output_file}")
    except Exception as e:
        print(f"保存 Word 文件失败：{e}")
    pdf_doc.close()

def process_pdf_files_in_directory(directory='file'):
    """处理指定目录下的所有PDF文件"""
    if not os.path.exists(directory):
        print(f"目录 {directory} 不存在")
        return

    pdf_files = [f for f in os.listdir(directory) if f.lower().endswith('.pdf')]
    if not pdf_files:
        print(f"在 {directory} 目录中没有找到PDF文件")
        return

    print(f"找到以下PDF文件：")
    for pdf_file in pdf_files:
        print(f"- {pdf_file}")

    for pdf_file in pdf_files:
        input_file = os.path.join(directory, pdf_file)
        output_file = os.path.join(directory, 
            f"{os.path.splitext(pdf_file)[0]}_translated.docx")
        print(f"\n开始处理：{pdf_file}")
        translate_pdf_to_word_with_styles_improved(input_file, output_file)

if __name__ == '__main__':
    import os
    process_pdf_files_in_directory()
